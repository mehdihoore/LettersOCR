import os
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
from pdf2image import convert_from_path
from docx import Document
from tkinter import Tk, Label, Button, filedialog
import multiprocessing
import tempfile
import arabic_reshaper
from bidi.algorithm import get_display
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Function to check if a corresponding docx file exists


def docx_exists(pdf_path):
    docx_path = os.path.splitext(pdf_path)[0] + '.docx'
    return os.path.exists(docx_path)

# Function to preprocess image for better OCR accuracy


def preprocess_image(image):
    # Convert to grayscale
    image = image.convert('L')
    # Enhance contrast
    image = ImageEnhance.Contrast(image).enhance(2)
    # Apply a filter to sharpen the image
    image = image.filter(ImageFilter.SHARPEN)
    return image

# Function to convert PDF to Word document only if a corresponding docx file doesn't exist


def convert_pdf_to_word(pdf_path):
    if not docx_exists(pdf_path):
        try:
            # Convert the PDF to images
            images = convert_from_path(pdf_path)

            # Create a new Word document
            doc = Document()

            # Loop through each image and extract text using OCR
            for i, image in enumerate(images):
                with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as temp_file:
                    image_path = temp_file.name
                    image = preprocess_image(image)
                    image.save(image_path, 'JPEG')

                # Extract text using OCR
                custom_config = r'--oem 3 --psm 6 -l fas+equ'
                text = pytesseract.image_to_string(image, config=custom_config)

                # Re-shape and display Persian text correctly
                reshaped_text = arabic_reshaper.reshape(text)
                bidi_text = get_display(reshaped_text)

                # Add the extracted text to the Word document
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(bidi_text)
                font = run.font
                font.name = 'B Nazanin'
                font.size = Pt(12)

                # Remove the temporary file
                os.remove(image_path)

            # Save the Word document with the same name as the PDF file
            docx_path = os.path.splitext(pdf_path)[0] + '.docx'
            doc.save(docx_path)
        except Exception as e:
            print(f"Error processing {pdf_path}: {e}")

# Function to handle folder selection and OCR on all PDF files in the folder


def select_folder_and_convert():
    folder_path = filedialog.askdirectory()
    if folder_path:
        pdf_files = [os.path.join(root, file) for root, _, files in os.walk(
            folder_path) for file in files if file.endswith('.pdf')]
        with multiprocessing.Pool() as pool:
            pool.map(convert_pdf_to_word, pdf_files)
        result_label.config(text="Conversion complete.")

# Create UI


def create_ui():
    root = Tk()
    root.title("PDF to Docx Converter")

    # UI elements
    label = Label(
        root, text="Select a folder containing PDF files to convert:")
    label.pack(pady=10)

    convert_button = Button(
        root, text="Select Folder and Convert", command=select_folder_and_convert)
    convert_button.pack(pady=10)

    global result_label
    result_label = Label(root, text="")
    result_label.pack(pady=10)

    # Run the UI loop
    root.mainloop()


if __name__ == "__main__":
    create_ui()
