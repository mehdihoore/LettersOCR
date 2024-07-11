import os
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import multiprocessing
import tempfile
import cv2
import numpy as np
from threading import Thread
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import multiprocessing
import tempfile
import cv2
import numpy as np
from threading import Thread


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def docx_exists(pdf_path):
    docx_path = os.path.splitext(pdf_path)[0] + '.docx'
    return os.path.exists(docx_path)


def preprocess_image(image):
    image = image.convert('L')
    image = ImageEnhance.Contrast(image).enhance(2)
    image = image.filter(ImageFilter.SHARPEN)
    return image


def detect_and_process_tables(image):
    cv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)
    thresh = cv2.threshold(
        gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]

    horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
    vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
    horizontal_lines = cv2.morphologyEx(
        thresh, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
    vertical_lines = cv2.morphologyEx(
        thresh, cv2.MORPH_OPEN, vertical_kernel, iterations=2)

    table_mask = cv2.bitwise_or(horizontal_lines, vertical_lines)
    contours, _ = cv2.findContours(
        table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    table_contours = [cnt for cnt in contours if cv2.contourArea(cnt) > 1000]

    return table_contours


def convert_pdf_to_word(pdf_path, language, save_dir, progress_callback=None):
    output_path = os.path.join(save_dir, os.path.basename(
        os.path.splitext(pdf_path)[0]) + '.docx')
    if not os.path.exists(output_path):
        try:
            images = convert_from_path(pdf_path)
            doc = Document()

            style = doc.styles['Normal']
            style.font.name = 'Arial' if language != 'fas' else 'B Nazanin'
            style.font.size = Pt(11)

            total_images = len(images)
            for i, image in enumerate(images):
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as temp_file:
                        image_path = temp_file.name
                        image = preprocess_image(image)
                        image.save(image_path, 'JPEG')

                    custom_config = r'--oem 3 --psm 6'
                    if language == 'fas':
                        custom_config += ' -l fas+ara+equ'
                    elif language == 'eng':
                        custom_config += ' -l eng'
                    elif language == 'deu':
                        custom_config += ' -l deu'
                    elif language == 'math':
                        custom_config += ' -l eng+equ'

                    text = pytesseract.image_to_string(
                        image, config=custom_config)

                    # Split text into paragraphs and add them to the document
                    paragraphs = text.split('\n\n')
                    for para_text in paragraphs:
                        paragraph = doc.add_paragraph(para_text.strip())
                        if language == 'fas':
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        else:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                    os.remove(image_path)

                    if progress_callback:
                        progress_callback((i + 1) / total_images * 100)
                except Exception as e:
                    print(f"Error processing page {
                          i+1} of {pdf_path}: {str(e)}")

            doc.save(output_path)

            # Verify the docx file
            try:
                Document(output_path)
                return True
            except Exception as e:
                print(f"Error verifying the created document {
                      output_path}: {str(e)}")
                return False
        except Exception as e:
            print(f"Error processing {pdf_path}: {str(e)}")
            return False
    return True



class SabaatPDFOCR(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("Sabaat PDF OCR")
        self.geometry("500x450")
        self.configure(bg="#E0E0E0")  # Light gray background

        # Set icon
        # Replace with your icon path
        self.iconbitmap(r"F:\Code\LettersOCR\icon BLACK.ico")

        self.create_widgets()

    def create_widgets(self):
        style = ttk.Style()
        style.theme_use('clam')
        style.configure('TFrame', background='#E0E0E0')
        style.configure('TLabel', background='#E0E0E0', foreground='#000000')
        style.configure('TButton', background='#FFFFFF', foreground='#000000')
        style.configure('TRadiobutton', background='#E0E0E0',
                        foreground='#000000')
        style.configure('TEntry', fieldbackground='#FFFFFF',
                        foreground='#000000')

        main_frame = ttk.Frame(self, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        title_label = ttk.Label(
            main_frame, text="Sabaat PDF OCR", font=("Helvetica", 16, "bold"))
        title_label.pack(pady=(0, 20))

        # Language selection
        lang_frame = ttk.Frame(main_frame)
        lang_frame.pack(fill=tk.X, pady=5)
        ttk.Label(lang_frame, text="Language:").pack(side=tk.LEFT)
        self.lang_var = tk.StringVar(value="fas")
        langs = [("Farsi", "fas"), ("English", "eng"),
                 ("German", "deu"), ("Math", "math")]
        for text, value in langs:
            ttk.Radiobutton(
                lang_frame, text=text, variable=self.lang_var, value=value).pack(side=tk.LEFT)

        # File or folder selection
        self.file_type_var = tk.StringVar(value="folder")
        file_frame = ttk.Frame(main_frame)
        file_frame.pack(fill=tk.X, pady=5)
        ttk.Radiobutton(file_frame, text="Single File",
                        variable=self.file_type_var, value="file").pack(side=tk.LEFT)
        ttk.Radiobutton(file_frame, text="Folder",
                        variable=self.file_type_var, value="folder").pack(side=tk.LEFT)

        # Input selection
        input_frame = ttk.Frame(main_frame)
        input_frame.pack(fill=tk.X, pady=5)
        self.input_entry = ttk.Entry(input_frame)
        self.input_entry.pack(side=tk.LEFT, expand=True, fill=tk.X)
        ttk.Button(input_frame, text="Browse",
                   command=self.browse_input).pack(side=tk.LEFT)

        # Output directory selection
        output_frame = ttk.Frame(main_frame)
        output_frame.pack(fill=tk.X, pady=5)
        ttk.Label(output_frame, text="Save to:").pack(side=tk.LEFT)
        self.output_entry = ttk.Entry(output_frame)
        self.output_entry.pack(side=tk.LEFT, expand=True, fill=tk.X)
        ttk.Button(output_frame, text="Browse",
                   command=self.browse_output).pack(side=tk.LEFT)

        # Convert button
        convert_button = ttk.Button(
            main_frame, text="Convert", command=self.start_conversion)
        convert_button.pack(fill=tk.X, pady=10)

        # Progress bar
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            main_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X, pady=10)

        self.status_label = ttk.Label(main_frame, text="")
        self.status_label.pack(pady=10)

        # Website address at bottom left
        website_label = tk.Label(self, text="@ www.sabaat.ir",
                                 fg="#000000", bg="#E0E0E0", font=("Helvetica", 8))
        website_label.pack(side=tk.BOTTOM, anchor=tk.SW, padx=5, pady=5)

    def browse_input(self):
        if self.file_type_var.get() == "file":
            path = filedialog.askopenfilename(
                filetypes=[("PDF files", "*.pdf")])
        else:
            path = filedialog.askdirectory()
        if path:
            self.input_entry.delete(0, tk.END)
            self.input_entry.insert(0, path)

    def browse_output(self):
        path = filedialog.askdirectory()
        if path:
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, path)

    def start_conversion(self):
        input_path = self.input_entry.get()
        output_path = self.output_entry.get()
        language = self.lang_var.get()

        if not input_path or not output_path:
            messagebox.showerror(
                "Error", "Please select input and output paths.")
            return

        if self.file_type_var.get() == "file":
            pdf_files = [input_path]
        else:
            pdf_files = [os.path.join(input_path, f) for f in os.listdir(
                input_path) if f.endswith('.pdf')]

        if not pdf_files:
            messagebox.showinfo(
                "No PDFs Found", "No PDF files were found in the selected location.")
            return

        self.progress_var.set(0)
        self.status_label.config(text="Converting...")

    

        Thread(target=self.process_files, args=(
            pdf_files, language, output_path)).start()

    def process_files(self, pdf_files, language, save_dir):
        total_files = len(pdf_files)
        converted_files = 0
        failed_files = []

        for pdf_file in pdf_files:
            try:
                success = convert_pdf_to_word(
                    pdf_file, language, save_dir, self.update_progress)
                if success:
                    converted_files += 1
                else:
                    failed_files.append(pdf_file)
            except Exception as e:
                print(f"Error processing {pdf_file}: {str(e)}")
                failed_files.append(pdf_file)

            self.progress_var.set((converted_files / total_files) * 100)
            self.update_idletasks()

        if failed_files:
            error_message = "The following files failed to convert or verify:\n" + \
                "\n".join(failed_files)
            messagebox.showerror("Conversion Errors", error_message)

        self.status_label.config(text=f"Conversion complete. {
                                converted_files}/{total_files} files converted successfully.")

        # Try to open one of the successfully converted files
        if converted_files > 0:
            try:
                successfully_converted = [
                    f for f in pdf_files if f not in failed_files][0]
                docx_path = os.path.join(save_dir, os.path.basename(
                    os.path.splitext(successfully_converted)[0]) + '.docx')
                os.startfile(docx_path)
            except Exception as e:
                print(f"Error opening converted file: {str(e)}")
                messagebox.showerror(
                    "Error", "Unable to open converted file. Please check the output directory manually.")

        # Open the output directory
        os.startfile(save_dir)

        # Open the folder containing the converted files
        if converted_files > 0:
            os.startfile(save_dir)

    def update_progress(self, value):
        self.progress_var.set(value)
        self.update_idletasks()


if __name__ == "__main__":
    app = SabaatPDFOCR()
    app.mainloop()
