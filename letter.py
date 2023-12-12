import os
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt
import re
import xlsxwriter
from persiantools.jdatetime import JalaliDate
import fitz
from fitz import Tools
# Function to convert PDF to Word document


def convert_pdf_to_word(pdf_path):
    # Convert the PDF to images
    images = convert_from_path(pdf_path)

    # Create a new Word document
    doc = Document()
    paragraph_style = doc.styles['Normal']
    paragraph_style.font.size = Pt(12)
    paragraph_style.font.name = 'Times New Roman'

    # Loop through each image and extract text using OCR
    for i, image in enumerate(images):
        # Save the image as a temporary file
        image_path = f'temp_image_{i}.jpg'
        image.save(image_path, 'JPEG')

        # Extract text using OCR
        text = pytesseract.image_to_string(
            Image.open(image_path), lang='fas')

        # Add a new page to the Word document
        if i > 0:
            doc.add_page_break()

        # Add the extracted text to the Word document
        paragraph = doc.add_paragraph(text)
        paragraph.style = paragraph_style

    # Save the Word document with the same name as the PDF file
    docx_path = os.path.splitext(pdf_path)[0] + '.docx'
    doc.save(docx_path)

    return docx_path

# Function to convert Jalali date to Gregorian


def convert_to_gregorian(persian_date):
    try:
        year, month, day = map(int, persian_date.split('-'))
        gregorian_date = JalaliDate(year, month, day).to_gregorian()
        return gregorian_date.strftime("%d/%m/%Y")
    except Exception as e:
        print(f"Error converting date {persian_date} to Gregorian: {e}")
        return ''


# Set the path to the folder containing the PDF files
pdf_folder_path = input('Please enter the path to the PDF folder: ')

# List to store data
data = []

# Traverse through the directory and subdirectories
for root, dirs, files in os.walk(pdf_folder_path):
    for file in files:
        # Check if the file is a PDF file
        if file.endswith('.pdf'):
            pdf_path = os.path.join(root, file)
            print(f'Converting {pdf_path} to Word document...')
            # Convert PDF to Word
            docx_path = convert_pdf_to_word(pdf_path)

            # Extract information from Word document
            try:
                doc = Document(docx_path)
                subject = ''
                date_str = ''
                for para in doc.paragraphs:
                    if 'موضوع' in para.text:
                        subject = para.text.split('موضوع')[1][:30].strip()
                        subject = re.sub(
                            r'[\\.,#+[\](\)\\/:*?<>|]', '-', subject)
                        # Add PDF name to the beginning of the subject with a hyphen
                        subject = f"{os.path.splitext(file)[0]}-{subject}"

                    elif 'تاریخ' in para.text:
                        date_str = para.text.split('تاریخ')[1][:15].strip()
                        date_str = re.sub(
                            r'[\\.,-_#+[\](\)\\/:*?<>|]', '-', date_str)
                pdf_doc = fitz.open(pdf_path)
    
                for page_num in range(pdf_doc.page_count):
                    page = pdf_doc[page_num]
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                    img.show()
                

                # Wait for user to view and close 
                input("Press Enter after reviewing PDF...") 
                pdf_doc.close()
                # Confirm or enter data manually
                confirm_subject = input(
                    f"Extracted subject: {subject}. Enter new subject or press Enter to confirm: ")
                if confirm_subject:
                    subject = confirm_subject

                confirm_date = input(
                    f"Extracted date: {date_str}. Enter new date or press Enter to confirm: ")
                if confirm_date:
                    date_str = confirm_date

                # Convert Jalali date to Gregorian
                date_gregorian = convert_to_gregorian(date_str)

                data.append((subject, date_str, date_gregorian))

                # Rename the PDF file with the new subject
                new_pdf_name = f"{subject}.pdf"
                new_pdf_path = os.path.join(root, new_pdf_name)
                os.rename(pdf_path, new_pdf_path)

            except Exception as e:
                print(f"Error processing {pdf_path}: {e}")

# Create Excel file and add worksheet
excel_file_path = os.path.join(pdf_folder_path, 'exported_data.xlsx')
workbook = xlsxwriter.Workbook(excel_file_path)
worksheet = workbook.add_worksheet()

# Write headers to the worksheet
worksheet.write(0, 0, 'موضوع')
worksheet.write(0, 1, 'تاریخ')
worksheet.write(0, 2, 'تاریخ (Gregorian)')

# Write data to the worksheet
for row, item in enumerate(data, start=1):
    print(f"Processing item: {item}")
    worksheet.write(row, 0, item[0])  # موضوع
    worksheet.write(row, 1, item[1])  # تاریخ
    worksheet.write(row, 2, item[2])  # تاریخ (Gregorian)

workbook.close()

print(f'Exported data to {excel_file_path}')
