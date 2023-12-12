import os
import pytesseract
from PIL import Image
from docx.shared import Pt
from fitz import Tools
import re
import xlsxwriter
from docx import Document
from persiantools.jdatetime import JalaliDate
import fitz
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from pdf2image import convert_from_path


class PDFConverterApp:
    def __init__(self, master):
        self.master = master
        master.title("PDF Converter")

        self.label = tk.Label(master, text="Select PDF Folder:")
        self.label.pack()

        self.folder_path_entry = tk.Entry(master, width=50)
        self.folder_path_entry.pack()

        self.browse_button = tk.Button(
            master, text="Browse", command=self.browse_folder)
        self.browse_button.pack()

        self.convert_button = tk.Button(
            master, text="Convert PDFs", command=self.convert_pdfs)
        self.convert_button.pack()

    def browse_folder(self):
        folder_selected = filedialog.askdirectory()
        self.folder_path_entry.delete(0, tk.END)
        self.folder_path_entry.insert(0, folder_selected)

    def convert_pdfs(self):
        pdf_folder_path = self.folder_path_entry.get()
        
        if not pdf_folder_path:
            messagebox.showerror("Error", "Please select a PDF folder.")
            return

        data = []

        for root, dirs, files in os.walk(pdf_folder_path):
            for file in files:
                if file.endswith('.pdf'):
                    pdf_path = os.path.join(root, file)
                    print(f'Converting {pdf_path} to Word document...')
                    images = convert_from_path(pdf_path)
                    doc = Document()
                    paragraph_style = doc.styles['Normal']
                    paragraph_style.font.size = Pt(12)
                    paragraph_style.font.name = 'Times New Roman'
                    for i, image in enumerate(images):
                        # Save the image as a temporary file
                        image_path = f'temp_image_{i}.jpg'
                        image.save(image_path, 'JPEG')

                        # Extract text using OCR
                        text = pytesseract.image_to_string(
                            Image.open(image_path), lang='fas')

                        # Add a new page to the Word document
                        if i > 0:
                            doc.add_page_break()

                        # Add the extracted text to the Word document
                        paragraph = doc.add_paragraph(text)
                        paragraph.style = paragraph_style

                    try:
                        
                        subject = ''
                        date_str = ''
                        for para in doc.paragraphs:
                            if 'موضوع' in para.text:
                                subject = para.text.split(
                                    'موضوع')[1][:30].strip()
                                subject = re.sub(
                                    r'[\\.,#+[\](\)\\/:*?<>|]', '-', subject)
                                # Add PDF name to the beginning of the subject with a hyphen
                                subject = f"{os.path.splitext(file)[0]}-{subject}"

                            elif 'تاریخ' in para.text:
                                date_str = para.text.split(
                                    'تاریخ')[1][:15].strip()
                                date_str = re.sub(
                                    r'[\\.,-_#+[\](\)\\/:*?<>|]', '-', date_str)
                        pdf_doc = fitz.open(pdf_path)

                        for page_num in range(pdf_doc.page_count):
                            page = pdf_doc[page_num]
                            pix = page.get_pixmap()
                            img = Image.frombytes(
                                "RGB", (pix.width, pix.height), pix.samples)
                            img.show()

                        # Wait for user to view and close
                        confirm_subject = messagebox.askyesno("Subject Confirmation", f"Extracted subject: {subject}. Do you want to change it?")
                        if confirm_subject:
                            new_subject = simpledialog.askstring("New Subject", "Enter new subject:")
                            if new_subject:
                                subject = new_subject

                        confirm_date = messagebox.askyesno("Date Confirmation", f"Extracted date: {date_str}. Do you want to change it?")
                        if confirm_date:
                            new_date = simpledialog.askstring("New Date", "Enter new date:")
                            if new_date:
                                date_str = new_date

                        # Convert Jalali date to Gregorian
                        date_gregorian = self.convert_to_gregorian(date_str)

                        data.append((subject, date_str, date_gregorian))

                        # Rename the PDF file with the new subject
                        new_pdf_name = f"{subject}.pdf"
                        new_pdf_path = os.path.join(root, new_pdf_name)
                        os.rename(pdf_path, new_pdf_path)

                    except Exception as e:
                        print(f"Error processing {pdf_path}: {e}")

        # Create Excel file and add worksheet
        excel_file_path = os.path.join(pdf_folder_path, 'exported_data.xlsx')
        workbook = xlsxwriter.Workbook(excel_file_path)
        worksheet = workbook.add_worksheet()

        # Write headers to the worksheet
        worksheet.write(0, 0, 'موضوع')
        worksheet.write(0, 1, 'تاریخ')
        worksheet.write(0, 2, 'تاریخ (Gregorian)')

        # Write data to the worksheet
        for row, item in enumerate(data, start=1):
            print(f"Processing item: {item}")
            worksheet.write(row, 0, item[0])  # موضوع
            worksheet.write(row, 1, item[1])  # تاریخ
            worksheet.write(row, 2, item[2])  # تاریخ (Gregorian)

        workbook.close()

        messagebox.showinfo("Conversion Complete",
                            "PDFs converted and data exported successfully.")

    def convert_to_gregorian(self, persian_date):
        try:
            year, month, day = map(int, persian_date.split('-'))
            gregorian_date = JalaliDate(year, month, day).to_gregorian()
            return gregorian_date.strftime("%d/%m/%Y")
        except Exception as e:
            print(f"Error converting date {persian_date} to Gregorian: {e}")
            return ''


if __name__ == "__main__":
    root = tk.Tk()
    app = PDFConverterApp(root)
    root.mainloop()
